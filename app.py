"""
Qbix Centre — Complete Web Application
Runs on Railway. Serves both the public website and the management app.
Data stored in qbix_data.json (committed to Railway or on persistent volume).
"""

import json
import os
import secrets
import hashlib
import hmac
import smtplib
import ssl
import threading
import time
import pyotp
from datetime import datetime, timedelta
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from functools import wraps
from pathlib import Path

from flask import (Flask, render_template, request, jsonify, redirect,
                   url_for, session, flash, send_file, abort)
from itsdangerous import URLSafeTimedSerializer, BadSignature, SignatureExpired
from dotenv import load_dotenv

load_dotenv()

# ── App setup ────────────────────────────────────────────────────────────────
app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', secrets.token_hex(32))
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=30)

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE_DIR  = Path(__file__).parent
DATA_FILE = BASE_DIR / os.environ.get('DATA_FILE', 'qbix_data.json')
BACKUP_DIR = BASE_DIR / 'backups'

# ── Config from environment ───────────────────────────────────────────────────
ADMIN_USERNAME    = os.environ.get('ADMIN_USERNAME', 'admin')
ADMIN_EMAIL       = os.environ.get('ADMIN_EMAIL', 'qbixcentre@outlook.com')
ADMIN_PHONE       = os.environ.get('ADMIN_PHONE', '4787379107')
APP_URL           = os.environ.get('APP_URL', 'http://localhost:5000')
FROM_EMAIL        = os.environ.get('FROM_EMAIL', 'noreply@qbixcentre.com')
FROM_NAME         = os.environ.get('FROM_NAME', 'Qbix Centre')
ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
GA_MEASUREMENT_ID = os.environ.get('GA_MEASUREMENT_ID', '')

# TOTP secret for 2FA (generated once and stored in env)
TOTP_SECRET = os.environ.get('TOTP_SECRET', pyotp.random_base32())

# Serializer for signed tokens (onboarding links, booking tokens)
serializer = URLSafeTimedSerializer(app.secret_key)

# In-memory stores (fine for single-instance Railway deployment)
_pending_2fa   = {}   # session_id -> {code, expires, purpose}
_booking_tokens = {}  # token -> member_email
_onboard_tokens = {}  # token -> {name, email, expires}

# ── Default data ──────────────────────────────────────────────────────────────
DEFAULT_DATA = {
    "offices": [
        {"id":"o1","num":"11","status":"Occupied","member":"HighBar Accounting","tenantStart":"9/1/2025","sqft":103,"dormer":None,"listDues":None},
        {"id":"o2","num":"12","status":"Occupied","member":"Pinnacle Accounting","tenantStart":"9/1/2025","sqft":139,"dormer":None,"listDues":None},
        {"id":"o3","num":"13","status":"Occupied","member":"Retail 1","tenantStart":"12/11/2018","sqft":196,"dormer":None,"listDues":None},
        {"id":"o4","num":"14","status":"Occupied","member":"Pinnacle Accounting","tenantStart":"9/1/2025","sqft":147,"dormer":None,"listDues":None},
        {"id":"o5","num":"15","status":"Occupied","member":"Pinnacle Accounting","tenantStart":"9/1/2025","sqft":150,"dormer":None,"listDues":None},
        {"id":"o6","num":"16","status":"Occupied","member":"Biren Patel Engineering","tenantStart":"6/8/2019","sqft":245,"dormer":None,"listDues":None},
        {"id":"o7","num":"17","status":"Occupied","member":"Biren Patel Engineering","tenantStart":"6/8/2019","sqft":196,"dormer":None,"listDues":None},
        {"id":"o8","num":"18","status":"Occupied","member":"Pettis Group","tenantStart":"12/1/2022","sqft":209,"dormer":None,"listDues":None},
        {"id":"o9","num":"19","status":"Occupied","member":"HighBar Accounting","tenantStart":"9/1/2025","sqft":176,"dormer":None,"listDues":None},
        {"id":"o10","num":"19A","status":"Occupied","member":"Preferred Provider Network","tenantStart":"9/2/2025","sqft":140,"dormer":None,"listDues":None},
        {"id":"o11","num":"19B","status":"Vacant","member":"","tenantStart":"","sqft":140,"dormer":None,"listDues":725},
        {"id":"o12","num":"21","status":"Occupied","member":"Gilbert Gomez CPA","tenantStart":"9/2/2025","sqft":90,"dormer":31,"listDues":None},
        {"id":"o13","num":"22","status":"Occupied","member":"McLendon Law","tenantStart":"10/1/2021","sqft":90,"dormer":31,"listDues":None},
        {"id":"o14","num":"23","status":"Occupied","member":"NAG Enterprise Group, LLC","tenantStart":"6/1/2025","sqft":99,"dormer":None,"listDues":None},
        {"id":"o15","num":"24","status":"Occupied","member":"HTNB Corp","tenantStart":"5/10/2025","sqft":90,"dormer":31,"listDues":None},
        {"id":"o16","num":"25","status":"Occupied","member":"McLendon Law","tenantStart":"10/1/2021","sqft":87,"dormer":None,"listDues":None},
        {"id":"o17","num":"26","status":"Vacant","member":"","tenantStart":"","sqft":87,"dormer":None,"listDues":500},
        {"id":"o18","num":"27","status":"Occupied","member":"Care Forth","tenantStart":"5/1/2021","sqft":87,"dormer":None,"listDues":None},
        {"id":"o19","num":"28","status":"Occupied","member":"Joshua David Nicholson","tenantStart":"10/1/2022","sqft":87,"dormer":None,"listDues":None},
        {"id":"o20","num":"29A","status":"Occupied","member":"Larry Fouche","tenantStart":"10/1/2025","sqft":142,"dormer":None,"listDues":None},
        {"id":"o21","num":"29B","status":"Vacant","member":"","tenantStart":"","sqft":158,"dormer":None,"listDues":725},
        {"id":"o22","num":"31","status":"Occupied","member":"Preferred Provider Network","tenantStart":"9/2/2025","sqft":139,"dormer":None,"listDues":None},
        {"id":"o23","num":"32","status":"Occupied","member":"National Youth Advocate Program","tenantStart":"5/1/2025","sqft":144,"dormer":None,"listDues":None},
        {"id":"o24","num":"33","status":"Occupied","member":"Wilson PC","tenantStart":"7/1/2023","sqft":161,"dormer":None,"listDues":None},
        {"id":"o25","num":"34","status":"Occupied","member":"Ram Bay","tenantStart":"10/1/2024","sqft":189,"dormer":None,"listDues":None},
        {"id":"o26","num":"35","status":"Occupied","member":"Rid A Critter","tenantStart":"2/1/2026","sqft":81,"dormer":None,"listDues":None},
        {"id":"o27","num":"36","status":"Occupied","member":"Ram Bay","tenantStart":"10/1/2024","sqft":136,"dormer":None,"listDues":None},
    ],
    "members": [],
    "occupants": [],
    "waitlist": [],
    "bookings": [],
    "templates": [
        {"id":"t1","name":"Power Outage","subject":"Power Outage Notice — Qbix Centre","body":"Dear {name},\n\nPlease be advised that Qbix Centre is currently experiencing a power outage. We are working to restore power as quickly as possible.\n\nWe apologize for any inconvenience and will keep you updated.\n\nThank you for your patience,\nQbix Centre Management"},
        {"id":"t2","name":"Monthly Dues Reminder","subject":"Monthly Dues Reminder — Qbix Centre","body":"Dear {name},\n\nThis is a friendly reminder that your monthly dues of {dues} are due. Please arrange payment at your earliest convenience.\n\nThank you,\nQbix Centre Management"},
        {"id":"t3","name":"Building Maintenance","subject":"Planned Maintenance Notice — Qbix Centre","body":"Dear {name},\n\nWe wanted to notify you that Qbix Centre will be undergoing scheduled maintenance. During this time, some services may be temporarily unavailable.\n\nBest regards,\nQbix Centre Management"},
        {"id":"t4","name":"General Notice","subject":"Important Notice from Qbix Centre","body":"Dear {name},\n\n[Your message here]\n\nBest regards,\nQbix Centre Management"},
    ],
    "lastBackup": "",
    "newsletter": [],
}


# ── Data helpers ──────────────────────────────────────────────────────────────
_data_lock = threading.Lock()

def load_data():
    if DATA_FILE.exists():
        with open(DATA_FILE, 'r', encoding='utf-8') as f:
            d = json.load(f)
        # Migrate: ensure new fields exist
        d.setdefault('bookings', [])
        d.setdefault('newsletter', [])
        for m in d.get('members', []):
            m.setdefault('attachments', [])
            m.setdefault('discount', 0)
            m.setdefault('agreementSent', '')
            m.setdefault('agreementSigned', '')
        for p in d.get('occupants', []):
            p.setdefault('dlAttachment', None)
        return d
    return json.loads(json.dumps(DEFAULT_DATA))

def save_data(data):
    with _data_lock:
        with open(DATA_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

def get_db():
    return load_data()

def net_dues(member):
    return max(0, (member.get('dues') or 0) - (member.get('discount') or 0))

def offices_for(data, name):
    return [o['num'] for o in data['offices'] if o.get('member') == name]

def hours_included(data, member_name):
    """Members get 6 hours per office they hold."""
    return len(offices_for(data, member_name)) * 6


# ── Email helper ──────────────────────────────────────────────────────────────
def send_email(to_email, to_name, subject, html_body, text_body=None):
    """Send email via SMTP. Configure SMTP_* env vars."""
    smtp_host = os.environ.get('SMTP_HOST', 'smtp.gmail.com')
    smtp_port = int(os.environ.get('SMTP_PORT', 587))
    smtp_user = os.environ.get('SMTP_USER', '')
    smtp_pass = os.environ.get('SMTP_PASS', '')

    if not smtp_user:
        print(f"[EMAIL] Would send to {to_email}: {subject}")
        return True

    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From']    = f"{FROM_NAME} <{FROM_EMAIL}>"
    msg['To']      = f"{to_name} <{to_email}>"

    if text_body:
        msg.attach(MIMEText(text_body, 'plain'))
    msg.attach(MIMEText(html_body, 'html'))

    try:
        context = ssl.create_default_context()
        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.starttls(context=context)
            server.login(smtp_user, smtp_pass)
            server.sendmail(FROM_EMAIL, to_email, msg.as_string())
        return True
    except Exception as e:
        print(f"[EMAIL ERROR] {e}")
        return False


def send_sms_code(phone, code):
    """Send SMS via email-to-SMS gateway (free) or Twilio if configured."""
    # Email-to-SMS gateways (free, carrier dependent)
    gateways = {
        'att':      f'{phone}@txt.att.net',
        'verizon':  f'{phone}@vtext.com',
        'tmobile':  f'{phone}@tmomail.net',
        'sprint':   f'{phone}@messaging.sprintpcs.com',
    }
    # Default: try AT&T gateway (most common in GA)
    carrier = os.environ.get('ADMIN_CARRIER', 'att')
    sms_email = gateways.get(carrier, gateways['att'])

    msg = MIMEText(f"Qbix Centre login code: {code}")
    msg['Subject'] = ''
    msg['From']    = FROM_EMAIL
    msg['To']      = sms_email

    smtp_host = os.environ.get('SMTP_HOST', 'smtp.gmail.com')
    smtp_port = int(os.environ.get('SMTP_PORT', 587))
    smtp_user = os.environ.get('SMTP_USER', '')
    smtp_pass = os.environ.get('SMTP_PASS', '')

    if not smtp_user:
        print(f"[SMS] Would send code {code} to {phone}")
        return True

    try:
        context = ssl.create_default_context()
        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.starttls(context=context)
            server.login(smtp_user, smtp_pass)
            server.sendmail(FROM_EMAIL, sms_email, msg.as_string())
        return True
    except Exception as e:
        print(f"[SMS ERROR] {e}")
        return False


# ── Auth helpers ──────────────────────────────────────────────────────────────
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def check_password(password):
    stored = os.environ.get('ADMIN_PASSWORD_HASH', '')
    if not stored:
        # First run — accept any password and prompt setup
        return True
    return hmac.compare_digest(hash_password(password), stored)

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('admin_authenticated'):
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    return decorated

def generate_code():
    return str(secrets.randbelow(900000) + 100000)  # 6-digit


# ── Template context ─────────────────────────────────────────────────────────
@app.context_processor
def inject_globals():
    return {'now': datetime.now()}


# ── Health check ──────────────────────────────────────────────────────────────
@app.route('/health')
def health():
    return jsonify({'ok': True, 'time': datetime.now().isoformat()})


# ══════════════════════════════════════════════════════════════════════════════
# PUBLIC WEBSITE ROUTES
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/')
def home():
    data = get_db()
    occupied = len([o for o in data['offices'] if o['status'] == 'Occupied'])
    total    = len(data['offices'])
    vacant   = [o for o in data['offices'] if o['status'] == 'Vacant']
    return render_template('public/home.html',
        occupied=occupied, total=total, vacant=vacant,
        ga_id=GA_MEASUREMENT_ID)

@app.route('/offices')
def offices_page():
    data = get_db()
    vacant = [o for o in data['offices'] if o['status'] == 'Vacant']
    return render_template('public/offices.html', vacant=vacant, ga_id=GA_MEASUREMENT_ID)

@app.route('/amenities')
def amenities():
    return render_template('public/amenities.html', ga_id=GA_MEASUREMENT_ID)

@app.route('/contact')
def contact():
    return render_template('public/contact.html', ga_id=GA_MEASUREMENT_ID)

@app.route('/contact', methods=['POST'])
def contact_submit():
    name    = request.form.get('name', '')
    email   = request.form.get('email', '')
    phone   = request.form.get('phone', '')
    message = request.form.get('message', '')
    # Email to admin
    send_email(
        ADMIN_EMAIL, 'Qbix Centre Admin',
        f'Website enquiry from {name}',
        f'<p><b>Name:</b> {name}<br><b>Email:</b> {email}<br><b>Phone:</b> {phone}</p><p>{message}</p>',
        f'Name: {name}\nEmail: {email}\nPhone: {phone}\n\n{message}'
    )
    flash('Thank you! We will be in touch shortly.', 'success')
    return redirect(url_for('contact'))

@app.route('/news')
def news():
    data = get_db()
    posts = sorted(data.get('newsletter', []), key=lambda x: x.get('date',''), reverse=True)
    return render_template('public/news.html', posts=posts, ga_id=GA_MEASUREMENT_ID)

@app.route('/news/<post_id>')
def news_post(post_id):
    data = get_db()
    post = next((p for p in data.get('newsletter', []) if p['id'] == post_id), None)
    if not post:
        abort(404)
    return render_template('public/news_post.html', post=post, ga_id=GA_MEASUREMENT_ID)


# ══════════════════════════════════════════════════════════════════════════════
# ONBOARDING FLOW (public — for prospective members)
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/onboard')
def onboard_home():
    """Landing page linked from website — general interest form."""
    return render_template('public/onboard_home.html', ga_id=GA_MEASUREMENT_ID)

@app.route('/onboard/<token>')
def onboard(token):
    """Personalized onboarding link sent by admin."""
    if token not in _onboard_tokens:
        return render_template('public/onboard_expired.html')
    info = _onboard_tokens[token]
    if datetime.now() > info['expires']:
        del _onboard_tokens[token]
        return render_template('public/onboard_expired.html')
    return render_template('public/onboard.html', token=token, info=info, ga_id=GA_MEASUREMENT_ID)

@app.route('/onboard/<token>/submit', methods=['POST'])
def onboard_submit(token):
    if token not in _onboard_tokens:
        return jsonify({'ok': False, 'error': 'Link expired'}), 400

    data = get_db()
    form = request.form

    # Create pending member
    member_id = '_' + secrets.token_hex(4)
    member = {
        'id':        member_id,
        'name':      form.get('company') or form.get('firstName') + ' ' + form.get('lastName'),
        'status':    'Pending',
        'start':     form.get('startDate', ''),
        'end':       '',
        'dues':      0,
        'discount':  0,
        'deposit':   0,
        'notes':     form.get('notes', ''),
        'email':     form.get('email', ''),
        'phone':     form.get('phone', ''),
        'address':   form.get('address', ''),
        'city':      form.get('city', ''),
        'state':     form.get('state', ''),
        'zip':       form.get('zip', ''),
        'website':   form.get('website', ''),
        'attachments':      [],
        'agreementSent':    '',
        'agreementSigned':  '',
        'emergencyName':    form.get('emergencyName', ''),
        'emergencyPhone':   form.get('emergencyPhone', ''),
        'emergencyRel':     form.get('emergencyRel', ''),
        'onboardedAt':      datetime.now().isoformat(),
    }
    data['members'].append(member)

    # Create pending occupant
    occ_id = '_' + secrets.token_hex(4)
    occupant = {
        'id':          occ_id,
        'name':        form.get('firstName', '') + ' ' + form.get('lastName', ''),
        'company':     member['name'],
        'phone':       form.get('phone', ''),
        'email':       form.get('email', ''),
        'office':      '',
        'endDate':     '',
        'status':      'Pending',
        'dlAttachment': None,
    }
    data['occupants'].append(occupant)
    save_data(data)

    # Notify admin
    send_email(
        ADMIN_EMAIL, 'Qbix Centre Admin',
        f'New onboarding submission: {member["name"]}',
        f'<p>A new prospect has completed the onboarding form.</p>'
        f'<p><b>Name:</b> {occupant["name"]}<br>'
        f'<b>Company:</b> {member["name"]}<br>'
        f'<b>Email:</b> {member["email"]}<br>'
        f'<b>Phone:</b> {member["phone"]}</p>'
        f'<p>Log in to your Qbix Centre dashboard to review and activate.</p>'
    )

    del _onboard_tokens[token]
    return jsonify({'ok': True})


# ══════════════════════════════════════════════════════════════════════════════
# CONFERENCE ROOM BOOKING (member-facing)
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/book')
def book_home():
    return render_template('public/book_home.html', ga_id=GA_MEASUREMENT_ID)

@app.route('/book/request-code', methods=['POST'])
def book_request_code():
    email = request.json.get('email', '').strip().lower()
    data  = get_db()

    # Check if active member with this email
    member = next((m for m in data['members']
                   if m.get('email','').lower() == email
                   and m.get('status') == 'Active'), None)
    if not member:
        return jsonify({'ok': False, 'error': 'Email not found in our active member list.'})

    code = generate_code()
    token = secrets.token_urlsafe(32)
    _pending_2fa[token] = {
        'code':    code,
        'email':   email,
        'name':    member['name'],
        'expires': datetime.now() + timedelta(minutes=10),
    }

    send_email(
        email, member['name'],
        'Your Qbix Centre booking code',
        f'<p>Your conference room booking code is:</p>'
        f'<h1 style="letter-spacing:8px;color:#2563eb">{code}</h1>'
        f'<p>This code expires in 10 minutes.</p>',
        f'Your Qbix Centre booking code: {code}\nExpires in 10 minutes.'
    )

    return jsonify({'ok': True, 'token': token})

@app.route('/book/verify', methods=['POST'])
def book_verify():
    token = request.json.get('token', '')
    code  = request.json.get('code', '').strip()

    entry = _pending_2fa.get(token)
    if not entry:
        return jsonify({'ok': False, 'error': 'Invalid or expired session.'})
    if datetime.now() > entry['expires']:
        del _pending_2fa[token]
        return jsonify({'ok': False, 'error': 'Code expired. Please request a new one.'})
    if entry['code'] != code:
        return jsonify({'ok': False, 'error': 'Incorrect code.'})

    # Issue booking session token
    booking_token = secrets.token_urlsafe(32)
    _booking_tokens[booking_token] = {
        'email': entry['email'],
        'name':  entry['name'],
        'expires': datetime.now() + timedelta(hours=2),
    }
    del _pending_2fa[token]
    return jsonify({'ok': True, 'bookingToken': booking_token})

@app.route('/book/calendar')
def book_calendar():
    bt = request.args.get('token', '')
    entry = _booking_tokens.get(bt)
    if not entry or datetime.now() > entry['expires']:
        return redirect(url_for('book_home'))

    data = get_db()
    member = next((m for m in data['members']
                   if m.get('email','').lower() == entry['email'].lower()), None)
    if not member:
        return redirect(url_for('book_home'))

    included = hours_included(data, member['name'])
    return render_template('public/book_calendar.html',
        token=bt,
        member_name=entry['name'],
        included_hours=included,
        ga_id=GA_MEASUREMENT_ID)

@app.route('/book/slots')
def book_slots():
    """Return booked slots for a given month."""
    bt = request.args.get('token', '')
    if bt not in _booking_tokens:
        return jsonify({'ok': False}), 401

    year  = int(request.args.get('year',  datetime.now().year))
    month = int(request.args.get('month', datetime.now().month))
    data  = get_db()

    slots = [b for b in data.get('bookings', [])
             if b.get('year') == year and b.get('month') == month
             and b.get('status') != 'cancelled']
    return jsonify({'ok': True, 'slots': slots})

@app.route('/book/create', methods=['POST'])
def book_create():
    bt = request.json.get('token', '')
    entry = _booking_tokens.get(bt)
    if not entry or datetime.now() > entry['expires']:
        return jsonify({'ok': False, 'error': 'Session expired'}), 401

    data   = get_db()
    member = next((m for m in data['members']
                   if m.get('email','').lower() == entry['email'].lower()), None)
    if not member:
        return jsonify({'ok': False, 'error': 'Member not found'}), 400

    date_str   = request.json.get('date', '')    # YYYY-MM-DD
    start_time = request.json.get('start', '')   # HH:MM
    end_time   = request.json.get('end', '')     # HH:MM
    title      = request.json.get('title', 'Meeting')

    try:
        dt = datetime.strptime(date_str, '%Y-%m-%d')
    except ValueError:
        return jsonify({'ok': False, 'error': 'Invalid date'}), 400

    # Check for conflicts
    for b in data.get('bookings', []):
        if (b.get('date') == date_str
                and b.get('status') != 'cancelled'
                and not (end_time <= b.get('start','') or start_time >= b.get('end',''))):
            return jsonify({'ok': False, 'error': 'That time slot is already booked.'})

    booking_id = '_' + secrets.token_hex(4)
    booking = {
        'id':          booking_id,
        'memberName':  member['name'],
        'memberEmail': entry['email'],
        'date':        date_str,
        'year':        dt.year,
        'month':       dt.month,
        'start':       start_time,
        'end':         end_time,
        'title':       title,
        'status':      'confirmed',
        'createdAt':   datetime.now().isoformat(),
    }
    data.setdefault('bookings', []).append(booking)
    save_data(data)

    # Confirmation email to member
    send_email(
        entry['email'], entry['name'],
        f'Conference Room Booking Confirmed — {date_str}',
        f'<p>Hi {entry["name"]},</p>'
        f'<p>Your conference room booking is confirmed:</p>'
        f'<ul><li><b>Date:</b> {date_str}</li>'
        f'<li><b>Time:</b> {start_time} – {end_time}</li>'
        f'<li><b>Room:</b> Qbix Centre Conference Room</li></ul>'
        f'<p>We look forward to seeing you!</p>'
        f'<p style="color:#666;font-size:12px">500A Northside Crossing, Macon, GA 31210</p>',
    )

    # Schedule reminder (24h before) in background thread
    def send_reminder():
        try:
            booking_dt = datetime.strptime(f'{date_str} {start_time}', '%Y-%m-%d %H:%M')
            reminder_dt = booking_dt - timedelta(hours=24)
            wait = (reminder_dt - datetime.now()).total_seconds()
            if wait > 0:
                time.sleep(wait)
            send_email(
                entry['email'], entry['name'],
                f'Reminder: Conference Room Tomorrow at {start_time}',
                f'<p>Hi {entry["name"]},</p>'
                f'<p>Just a reminder that you have the Qbix Centre conference room booked tomorrow:</p>'
                f'<ul><li><b>Date:</b> {date_str}</li>'
                f'<li><b>Time:</b> {start_time} – {end_time}</li></ul>'
                f'<p>See you tomorrow!</p>',
            )
        except Exception as e:
            print(f'Reminder error: {e}')

    threading.Thread(target=send_reminder, daemon=True).start()

    return jsonify({'ok': True, 'booking': booking})

@app.route('/book/cancel', methods=['POST'])
def book_cancel():
    bt = request.json.get('token', '')
    entry = _booking_tokens.get(bt)
    if not entry or datetime.now() > entry['expires']:
        return jsonify({'ok': False}), 401

    booking_id = request.json.get('bookingId', '')
    data = get_db()
    booking = next((b for b in data.get('bookings', [])
                    if b['id'] == booking_id
                    and b['memberEmail'].lower() == entry['email'].lower()), None)
    if not booking:
        return jsonify({'ok': False, 'error': 'Booking not found'}), 404

    booking['status'] = 'cancelled'
    save_data(data)
    return jsonify({'ok': True})


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN LOGIN (2-factor)
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    # Check if already fully authenticated
    if session.get('admin_authenticated'):
        return redirect(url_for('admin_dashboard'))

    if request.method == 'POST':
        username = request.form.get('username', '')
        password = request.form.get('password', '')

        if username == ADMIN_USERNAME and check_password(password):
            # Password correct — send 2FA code
            code  = generate_code()
            sid   = secrets.token_urlsafe(16)
            _pending_2fa[sid] = {
                'code':    code,
                'expires': datetime.now() + timedelta(minutes=10),
                'purpose': 'admin',
            }
            session['admin_2fa_sid'] = sid

            # Send code via SMS
            send_sms_code(ADMIN_PHONE, code)

            return redirect(url_for('admin_2fa'))
        else:
            flash('Invalid username or password.', 'error')

    return render_template('admin/login.html')

@app.route('/admin/2fa', methods=['GET', 'POST'])
def admin_2fa():
    sid = session.get('admin_2fa_sid')
    if not sid or sid not in _pending_2fa:
        return redirect(url_for('admin_login'))

    if request.method == 'POST':
        code  = request.form.get('code', '').strip()
        entry = _pending_2fa.get(sid)

        if not entry or datetime.now() > entry['expires']:
            flash('Code expired. Please log in again.', 'error')
            return redirect(url_for('admin_login'))

        if entry['code'] != code:
            flash('Incorrect code. Please try again.', 'error')
            return render_template('admin/2fa.html')

        # Success — fully authenticated
        del _pending_2fa[sid]
        session.pop('admin_2fa_sid', None)
        session['admin_authenticated'] = True
        session['admin_login_time']    = datetime.now().isoformat()
        session.permanent = True

        return redirect(url_for('admin_dashboard'))

    return render_template('admin/2fa.html')

@app.route('/admin/logout')
def admin_logout():
    session.clear()
    return redirect(url_for('home'))


# ══════════════════════════════════════════════════════════════════════════════
# ADMIN — MANAGEMENT APP
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/admin')
@login_required
def admin_dashboard():
    data = get_db()
    active  = [m for m in data['members'] if m['status'] == 'Active']
    pending = [m for m in data['members'] if m['status'] == 'Pending']
    occ     = len([o for o in data['offices'] if o['status'] == 'Occupied'])
    vac     = len([o for o in data['offices'] if o['status'] == 'Vacant'])
    gross   = sum((m.get('dues') or 0) for m in active)
    net_rev = sum(net_dues(m) for m in active)
    dep     = sum((m.get('deposit') or 0) for m in active)
    endings = sorted([m for m in data['members'] if m.get('end')],
                     key=lambda m: m['end'])[:5]
    vacant  = [o for o in data['offices'] if o['status'] == 'Vacant']

    # Upcoming bookings (next 7 days)
    today = datetime.now().date()
    upcoming = sorted(
        [b for b in data.get('bookings', [])
         if b.get('status') == 'confirmed'
         and datetime.strptime(b['date'], '%Y-%m-%d').date() >= today],
        key=lambda b: (b['date'], b['start'])
    )[:10]

    return render_template('admin/dashboard.html',
        active=active, pending=pending,
        occ=occ, vac=vac, gross=gross, net_rev=net_rev, dep=dep,
        endings=endings, vacant=vacant, upcoming_bookings=upcoming,
        total=len(data['offices']),
        data=data)

# ── Data API (used by admin JS frontend) ─────────────────────────────────────

@app.route('/admin/api/data')
@login_required
def api_data():
    return jsonify(get_db())

@app.route('/admin/api/save', methods=['POST'])
@login_required
def api_save():
    try:
        data = request.json
        save_data(data)
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/admin/api/backup')
@login_required
def api_backup():
    try:
        BACKUP_DIR.mkdir(exist_ok=True)
        today = datetime.now().strftime('%Y-%m-%d')
        dest  = BACKUP_DIR / f'qbix-backup-{today}.json'
        import shutil
        shutil.copy2(DATA_FILE, dest)
        # Update lastBackup in data
        data = get_db()
        data['lastBackup'] = today
        save_data(data)
        return jsonify({'ok': True, 'path': str(dest)})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

# ── Onboarding link generator ─────────────────────────────────────────────────

@app.route('/admin/api/onboard-link', methods=['POST'])
@login_required
def generate_onboard_link():
    name  = request.json.get('name', '')
    email = request.json.get('email', '')
    token = secrets.token_urlsafe(16)
    _onboard_tokens[token] = {
        'name':    name,
        'email':   email,
        'expires': datetime.now() + timedelta(days=7),
    }
    link = f'{APP_URL}/onboard/{token}'

    # Email the link to the prospect
    if email:
        send_email(
            email, name,
            'Welcome to Qbix Centre — Complete Your Application',
            f'<p>Hi {name},</p>'
            f'<p>Thank you for your interest in Qbix Centre! Please click the link below to complete your membership application:</p>'
            f'<p><a href="{link}" style="background:#2563eb;color:#fff;padding:12px 24px;border-radius:6px;text-decoration:none;display:inline-block">Complete My Application</a></p>'
            f'<p>This link expires in 7 days.</p>'
            f'<p>If you have any questions, reply to this email or call (478) 787-0532.</p>'
            f'<p>We look forward to welcoming you!</p>'
            f'<p>— The Qbix Centre Team</p>',
        )

    return jsonify({'ok': True, 'link': link, 'token': token})

# ── Agreement generator ───────────────────────────────────────────────────────

@app.route('/admin/api/generate-agreement/<member_id>')
@login_required
def generate_agreement(member_id):
    """Generate a filled-in membership agreement as .docx"""
    data   = get_db()
    member = next((m for m in data['members'] if m['id'] == member_id), None)
    if not member:
        abort(404)

    offices = offices_for(data, member['name'])
    office_str = ', '.join(offices) if offices else 'TBD'
    dues_str   = f"${member.get('dues', 0):,}/month"
    deposit_str = f"${member.get('deposit', 0):,}"
    start_str  = member.get('start', '')
    today_str  = datetime.now().strftime('%B %d, %Y')

    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io

        doc = DocxDocument()

        # Title
        title = doc.add_heading('Qbix Centre Membership Agreement', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f'Date: {today_str}')
        doc.add_paragraph(f'Location: 500A Northside Crossing, Macon, GA 31210')
        doc.add_paragraph('')

        # Member info table
        doc.add_heading('Member Information', level=1)
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        rows_data = [
            ('Member / Company', member.get('name', '')),
            ('Contact Email', member.get('email', '')),
            ('Contact Phone', member.get('phone', '')),
            ('Office(s)', office_str),
            ('Monthly Dues', dues_str),
            ('Deposit', deposit_str),
        ]
        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

        doc.add_paragraph('')

        # Agreement text
        doc.add_heading('1. Membership & Fees', level=1)
        doc.add_paragraph(
            f'Member agrees to pay monthly dues of {dues_str} due on the 1st of each month '
            f'via auto-draft. A refundable deposit of {deposit_str} is required. '
            f'The initial term is six (6) months beginning {start_str}, automatically renewing '
            f'unless terminated with 30 days written notice prior to end of term.'
        )

        doc.add_heading('2. Access & Use', level=1)
        doc.add_paragraph(
            '24/7 access is provided via card key/fob and access code. Codes must remain '
            'confidential. Conference room use is limited to 6 included hours per office per month; '
            'overages billed at $50/hour. Workspace is for professional business use only.'
        )

        doc.add_heading('3. Conduct & Responsibilities', level=1)
        doc.add_paragraph(
            'Member agrees to maintain a professional, respectful environment. No hazardous '
            'materials, smoking, or pets. Member is responsible for their own actions and those '
            'of their guests. Damages must be reported and paid for promptly.'
        )

        doc.add_heading('4. Insurance & Liability', level=1)
        doc.add_paragraph(
            'Qbix Centre does not provide insurance for member property. Members should obtain '
            'their own coverage. RoseAn Properties, LLC and affiliates are not liable for theft, '
            'loss, or damage. Members indemnify Qbix Centre against claims arising from their use.'
        )

        doc.add_heading('5. Termination', level=1)
        doc.add_paragraph(
            'Member may terminate with 30 days written notice prior to end of term. '
            'Management may terminate immediately for rule violations without refund.'
        )

        doc.add_heading('6. General Provisions', level=1)
        doc.add_paragraph(
            'This agreement provides a license to use shared workspace — it is not a lease or '
            'tenancy. Rules may be updated at any time with reasonable notice. Member must have '
            'authority to bind their company to this agreement.'
        )

        # Signature block
        doc.add_paragraph('')
        doc.add_heading('Signatures', level=1)
        sig_table = doc.add_table(rows=4, cols=2)
        sig_table.style = 'Table Grid'
        sig_rows = [
            ('Member Signature', ''),
            ('Print Name', member.get('name', '')),
            ('Date', ''),
            ('Qbix Centre Representative', ''),
        ]
        for i, (label, value) in enumerate(sig_rows):
            sig_table.rows[i].cells[0].text = label
            sig_table.rows[i].cells[1].text = value

        # Save to buffer
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        filename = f"Qbix_Agreement_{member['name'].replace(' ','_')}_{today_str.replace(' ','_')}.docx"
        return send_file(buf, as_attachment=True,
                        download_name=filename,
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except ImportError:
        return jsonify({'ok': False, 'error': 'python-docx not installed'}), 500

# ── Newsletter generator (AI-assisted) ───────────────────────────────────────

@app.route('/admin/api/generate-newsletter', methods=['POST'])
@login_required
def generate_newsletter():
    if not ANTHROPIC_API_KEY:
        return jsonify({'ok': False, 'error': 'Anthropic API key not configured'}), 400

    data   = get_db()
    active = [m for m in data['members'] if m['status'] == 'Active']
    occ    = len([o for o in data['offices'] if o['status'] == 'Occupied'])
    vac    = len([o for o in data['offices'] if o['status'] == 'Vacant'])
    month  = datetime.now().strftime('%B %Y')

    context = (
        f"Qbix Centre is a professional coworking space in Macon, Georgia at 500A Northside Crossing. "
        f"Current stats: {occ} offices occupied, {vac} vacant, {len(active)} active members. "
        f"Month: {month}. "
        f"Amenities: 24/7 access, AT&T Fiber, Starbucks coffee, furnished offices, conference room, free parking."
    )

    custom_notes = request.json.get('notes', '')

    try:
        import urllib.request
        import json as json_mod

        payload = {
            'model': 'claude-sonnet-4-20250514',
            'max_tokens': 1000,
            'messages': [{
                'role': 'user',
                'content': (
                    f'Write a warm, friendly, professional monthly newsletter for Qbix Centre. '
                    f'Keep it concise — 3-4 short paragraphs. Include a welcoming opener, '
                    f'a community update, any relevant seasonal note, and a friendly closing. '
                    f'Context: {context}. '
                    f'Additional notes from manager: {custom_notes}. '
                    f'Format as HTML suitable for email. Use a warm, community-focused tone.'
                )
            }]
        }

        req = urllib.request.Request(
            'https://api.anthropic.com/v1/messages',
            data=json_mod.dumps(payload).encode(),
            headers={
                'Content-Type': 'application/json',
                'x-api-key': ANTHROPIC_API_KEY,
                'anthropic-version': '2023-06-01',
            }
        )
        with urllib.request.urlopen(req, timeout=30) as resp:
            result = json_mod.loads(resp.read())
            draft = result['content'][0]['text']
            return jsonify({'ok': True, 'draft': draft})

    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/admin/api/publish-newsletter', methods=['POST'])
@login_required
def publish_newsletter():
    """Save newsletter and optionally email to all active members."""
    data    = get_db()
    subject = request.json.get('subject', f'Qbix Centre Newsletter — {datetime.now().strftime("%B %Y")}')
    body    = request.json.get('body', '')
    send    = request.json.get('send', False)

    post_id = '_' + secrets.token_hex(4)
    post = {
        'id':      post_id,
        'subject': subject,
        'body':    body,
        'date':    datetime.now().isoformat(),
        'sent':    send,
    }
    data.setdefault('newsletter', []).append(post)
    save_data(data)

    if send:
        active_with_email = [m for m in data['members']
                             if m['status'] == 'Active' and m.get('email')]
        sent_count = 0
        for m in active_with_email:
            ok = send_email(m['email'], m['name'], subject, body)
            if ok:
                sent_count += 1
        return jsonify({'ok': True, 'sent': sent_count, 'postId': post_id})

    return jsonify({'ok': True, 'sent': 0, 'postId': post_id})

# ── Booking management (admin view) ──────────────────────────────────────────

@app.route('/admin/bookings')
@login_required
def admin_bookings():
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/api/booking-cancel', methods=['POST'])
@login_required
def admin_cancel_booking():
    booking_id = request.json.get('bookingId', '')
    data = get_db()
    booking = next((b for b in data.get('bookings', []) if b['id'] == booking_id), None)
    if not booking:
        return jsonify({'ok': False}), 404
    booking['status'] = 'cancelled'
    save_data(data)
    return jsonify({'ok': True})

# ── Monthly usage email (called by scheduler or manually) ────────────────────

@app.route('/admin/api/send-monthly-usage', methods=['POST'])
@login_required
def send_monthly_usage():
    data  = get_db()
    month = request.json.get('month', datetime.now().month)
    year  = request.json.get('year',  datetime.now().year)

    active_with_email = [m for m in data['members']
                         if m['status'] == 'Active' and m.get('email')]
    sent = 0
    month_name = datetime(year, month, 1).strftime('%B %Y')

    for member in active_with_email:
        # Count their bookings that month
        member_bookings = [
            b for b in data.get('bookings', [])
            if b.get('memberEmail','').lower() == member.get('email','').lower()
            and b.get('year') == year
            and b.get('month') == month
            and b.get('status') != 'cancelled'
        ]

        # Calculate hours used
        total_minutes = 0
        for b in member_bookings:
            try:
                s = datetime.strptime(b['start'], '%H:%M')
                e = datetime.strptime(b['end'],   '%H:%M')
                total_minutes += int((e - s).total_seconds() / 60)
            except Exception:
                pass
        hours_used     = round(total_minutes / 60, 1)
        included       = hours_included(data, member['name'])
        hours_remaining = max(0, included - hours_used)

        if hours_used == 0 and not member_bookings:
            # Skip members who didn't use the room
            continue

        booking_list = ''.join(
            f'<li>{b["date"]} {b["start"]}–{b["end"]}: {b.get("title","Meeting")}</li>'
            for b in member_bookings
        )

        send_email(
            member['email'], member['name'],
            f'Your Qbix Centre Conference Room Summary — {month_name}',
            f'<p>Hi {member["name"]},</p>'
            f'<p>Here\'s a cheerful summary of your conference room use in {month_name}! 🎉</p>'
            f'<ul>{booking_list}</ul>'
            f'<p><b>Total hours used:</b> {hours_used} of your {included} included hours</p>'
            f'<p><b>Hours remaining this month:</b> {hours_remaining}</p>'
            f'<p>Thank you for being part of the Qbix Centre community. See you next month!</p>'
            f'<p style="color:#666;font-size:12px">Questions? Reply to this email or call (478) 787-0532</p>',
        )
        sent += 1

    return jsonify({'ok': True, 'sent': sent})


# ── Setup route (first run only) ─────────────────────────────────────────────

@app.route('/admin/setup', methods=['GET', 'POST'])
def admin_setup():
    """First-run setup to set admin password."""
    if os.environ.get('ADMIN_PASSWORD_HASH'):
        return redirect(url_for('admin_login'))

    if request.method == 'POST':
        password = request.form.get('password', '')
        confirm  = request.form.get('confirm', '')
        if password != confirm:
            flash('Passwords do not match.', 'error')
        elif len(password) < 8:
            flash('Password must be at least 8 characters.', 'error')
        else:
            hashed = hash_password(password)
            flash(
                f'Setup complete! Add this to your Railway environment variables: '
                f'ADMIN_PASSWORD_HASH={hashed}',
                'success'
            )
    return render_template('admin/setup.html')


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8765))
    app.run(host='0.0.0.0', port=port, debug=False)
